# 导入必要的库
from urllib.parse import urlparse
import requests  # 导入requests库用于发送HTTP请求
from requests.exceptions import RequestException  # 用于异常处理
from collections import OrderedDict  # 有序字典，用于保持请求头的顺序
import hmac  # 用于HMAC加密算法
import hashlib  # 提供多种安全哈希算法
import base64  # 用于Base64编码
import time    # 动态时间戳
import json    # 用于解析JSON
import urllib3 # 用于控制SSL警告
import os      # 用于处理文件路径

# ===================== 全局配置参数 =====================
# 注意：以下参数需要从海康开放平台获取，请确保与平台配置完全一致
URL = "/artemis/api/pms/v1/crossRecords/page"  # 请求的API接口路径
appKey = "23003292"  # 应用唯一标识（类似账号）
appSecret = "huYd9V9zPd3aNlyc8CDN"  # 应用密钥（类似密码，需严格保密）
x_ca_signature_headers = "x-ca-key"  # 指定参与签名的请求头字段


# ===================== 签名工具类 =====================
class SignUtil:
    @staticmethod
    def sign(method, url, headers, app_secret):
        """
        生成海康API请求签名
        参数说明：
        method: HTTP请求方法（大写字母），例如：POST
        url: 请求的API路径，例如：/artemis/api/...
        headers: 包含请求头的字典
        app_secret: 应用密钥（用于加密的关键参数）
        """
        # 步骤1：构建待签名字符串（这是签名的核心内容）
        string_to_sign = (
            f"{method}\n"  # HTTP方法（例如：POST）
            f"{headers['Accept']}\n"  # 接受的数据类型（例如：*/* 表示接受所有类型）
            f"{headers['Content-Type']}\n"  # 发送的内容类型（例如：application/json）
            f"x-ca-key:{headers['X-Ca-Key']}\n"  # 强制小写的关键头（注意冒号后的值）
            # f"x-ca-timestamp:{headers['X-Ca-Timestamp']}\n"  # 时间戳（签名使用可能访问不了）
            f"{url}"  # API请求路径（注意最后没有换行符）
        )

        # 步骤2：将字符串转换为字节流（加密函数需要字节类型数据）
        byte_stream = string_to_sign.encode("utf-8")  # 必须使用UTF-8编码

        # 调试输出（开发时可查看签名内容）
        print("【待签名字符串】")
        print(repr(string_to_sign))  # 显示原始字符串（包含换行符）
        print(f"字节长度: {len(byte_stream)} ")  # 打印字节长度用于校验
        print("字节内容:", byte_stream)  # 显示实际参与加密的字节数据

        # 步骤3：使用HMAC-SHA256算法生成签名
        digest = hmac.new(
            app_secret.encode("utf-8"),  # 将密钥转换为字节
            byte_stream,  # 待签名字节流
            hashlib.sha256  # 指定哈希算法
        ).digest()  # 生成二进制签名

        # 步骤4：将二进制签名转换为Base64字符串
        return base64.b64encode(digest).decode()  # 解码为字符串格式


# ===================== 文档生成模块 =====================
class ReportGenerator:
    @staticmethod
    def save_to_md(data, filename="API报告.md"):
        """生成Markdown格式报告"""
        desktop = os.path.join(os.path.expanduser('~'), 'Desktop')
        path = os.path.join(desktop, filename)

        # 构建MD内容
        md_content = "# API请求报告\n\n"
        md_content += f"生成时间：{time.strftime('%Y-%m-%d %H:%M:%S')}\n\n"

        for item in data:
            md_content += f"## {item['name']}\n"
            md_content += f"- **接口地址**: `{item['url']}`\n"
            md_content += f"- **状态码**: {item['status_code']}\n"
            md_content += "### 响应内容\n```json\n"
            md_content += json.dumps(item['response'], indent=2, ensure_ascii=False)
            md_content += "\n```\n\n"
        # 写入文件
        with open(path, 'w', encoding='utf-8') as f:
            f.write(md_content)
        print(f"MD报告已保存至：{path}")

    @staticmethod
    def save_to_docx(data, filename="API报告.docx"):
        """生成Word格式报告（需要python-docx库）"""
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.style import WD_STYLE_TYPE
        except ImportError:
            print("请先安装python-docx库：pip install python-docx")
            return
        desktop = os.path.join(os.path.expanduser('~'), 'Desktop')
        path = os.path.join(desktop, filename)

        doc = Document()
        # 添加标题
        title = doc.add_heading('API请求报告', 0)
        title.add_run(f"\n生成时间：{time.strftime('%Y-%m-%d %H:%M:%S')}").font.size = Pt(12)
        # 修改表格样式设置方式
        table = doc.add_table(rows=2, cols=2)
        # 方法一：使用内置样式名称（推荐）
        table.style = 'Light List Accent 1'  # 有效的内置样式名称

        # 添加每个接口的信息
        for item in data:
            # 添加接口标题
            doc.add_heading(item['name'], level=1)

            # 创建信息表格
            table = doc.add_table(rows=2, cols=2)
            table.style = 'LightShading-Accent1'
            # 填充表格
            cells = table.rows[0].cells
            cells[0].text = "接口地址"
            cells[1].text = item['url']

            cells = table.rows[1].cells
            cells[0].text = "状态码"
            cells[1].text = str(item['status_code'])

            # 添加响应内容
            doc.add_heading('响应内容', level=2)
            content = json.dumps(item['response'], indent=2, ensure_ascii=False)
            doc.add_paragraph(content)

            doc.add_page_break()

        doc.save(path)
        print(f"Word报告已保存至：{path}")


# ===================== 主函数（使用示例） =====================
def main(request_url=None):
    # 禁用SSL警告（添加在requests导入之后）
    urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

    api_endpoints = [
        {
            "name": "获取停车库列表",
            "url": "https://172.19.9.9:443/artemis/api/resource/v1/park/parkList",
            "method": "POST",
            "payload": {"parkIndexCodes": ""}
        },
        {
            "name": "获取出入口列表",
            "url": "https://172.19.9.9:443/artemis/api/resource/v1/entrance/entranceList",
            "method": "POST",
            "payload": {
                "parkIndexCodes": "ec0d500279934dbd8a17a6e38963deea"  # 取停车库列表返回中的“parkIndexCode”
            }
        },
        {
                "name": "查询停车库剩余车位数",
                "url": "https://172.19.9.9:443/artemis/api/pms/v1/park/remainSpaceNum",
                "method": "POST",
                "payload": {
                    "parkSyscode": "ec0d500279934dbd8a17a6e38963deea"   # 取出入口列表中返回的“entranceIndexCode”

                }
        },
        {
                "name": "获取车道列表",
                "url": "https://172.19.9.9:443/artemis/api/resource/v1/roadway/roadwayList",
                "method": "POST",
                "payload": {
                    "entranceIndexCodes": "dd9d05054821446c84c563eed6ea6698"  # 取出入口列表中返回的"entranceIndexCode"
                }
        },
        {
                "name": "根据车道编码反控道闸",
                "url": "https://172.19.9.9:443/artemis/api/pms/v1/deviceControl",
                "method": "POST",
                "payload": {
                    "roadwaySyscode": "952a06dbf563412eaa45fa30bdc4e5fc", # 取车道列表中返回的“roadwayIndexCode”
                    #"command": 1       # 0关闸，1开闸，3常开
                }
        },
        {
            "name": "查询过车记录",
            "url": "https://172.19.9.9:443/artemis/api/pms/v1/crossRecords/page",
            "method": "POST",
            "payload": {
                "startTime": "2025-04-12T15:00:00+08:00",
                "endTime": "2025-04-15T15:00:00+08:00",
                "pageNo": 1,
                "pageSize": 1     # 1-1000
            }
        },
        {
            "name": "查询车辆抓拍图片",
            "url": "https://172.19.9.9:443/artemis/api/pms/v1/image",
            "method": "POST",
            "payload": {
                "aswSyscode": "hnj5h245h5234h45345y",
                "picUri": "/pic?=d7ei703i10cd*73a-d5108a--22cd0c9d6592aiid="
            }
        },
        {
            "name": "查询区域列表v2",
            "url": "https://172.19.9.9:443/artemis/api/irds/v2/region/nodesByParams",
            "method": "POST",
            "payload": {
                "resourceType": "snapshot",  # 附录A.3 资源类型/资源权限码；
                "pageNo": 1,
                "pageSize": 100
            }
        },
        {
            "name": "获取停车场设备在线状态",
            "url": "https://172.19.9.9:443/artemis/api/nms/v1/online/pms/device/get",
            "method": "POST",
            "payload": {
                "regionId": "root000000",     # 查询区域列表v2返回的"regionId"
                "resourceType": "snapshot",   # 附录A.88 停车场网管资源类型
                "pageNo": 1,
                "pageSize": 100
            }
        }
    ]
    results = []  # 存储所有接口的返回结果
    for api in api_endpoints:
        # 提取API路径（关键修改）
        parsed_url = urlparse(api['url'])
        api_path = parsed_url.path  # 获取/artemis开头的路径

        # 动态生成时间戳（重要！）
        headers = OrderedDict()
        headers["Accept"] = "*/*"
        headers["Content-Type"] = "application/json"
        headers["x-ca-signature-headers"] = x_ca_signature_headers
        headers["X-Ca-Key"] = appKey
        headers["X-Ca-Timestamp"] = str(int(time.time() * 1000))  # 动态时间戳

        # 生成当前API的签名（关键修改）
        sign_result = SignUtil.sign(
            method=api['method'],
            url=api_path,  # 使用当前API路径
            headers=headers,
            app_secret=appSecret
        )
        headers["x-ca-signature"] = sign_result

        # 发送POST请求（跳过SSL验证，生产环境不推荐）
        try:
            print(f"\n【正在请求接口】{api['name']}")
            response = requests.request(
                method=api['method'],
                url=api['url'],
                headers=headers,
                json=api['payload'],
                verify=False
            )
            response.raise_for_status()  # 自动抛出HTTP错误

            # 使用json库解析响应内容
            try:
                response_data = response.json()
                results.append({
                    "name": api['name'],
                    "url": api['url'],
                    "status_code": response.status_code,
                    "response": response_data
                })
                print(f"状态码：{response.status_code} 请求成功！")
            except json.JSONDecodeError:
                print("响应内容不是有效的JSON格式")
                results.append({
                    "name": api['name'],
                    "url": api['url'],
                    "status_code": response.status_code,
                    "response": response.text
                })
        except RequestException as e:
            print(f"请求失败：{str(e)}")
            results.append({
                "name": api['name'],
                "url": api['url'],
                "status_code": e.response.status_code if e.response else 500,
                "response": str(e)
            })

    # 生成报告
    if results:
        ReportGenerator.save_to_md(results)
        ReportGenerator.save_to_docx(results)
    else:
        print("没有可保存的结果")


if __name__ == "__main__":
    main()
